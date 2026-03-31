from flask import Flask, render_template, request, redirect, url_for, flash, session
from werkzeug.security import generate_password_hash, check_password_hash
import sqlite3
import pymysql
from pathlib import Path
from datetime import datetime
import pandas as pd
import os
import sys
from werkzeug.utils import secure_filename
from flask import send_file, abort
from docx import Document
#import comtypes.client
import uuid
#import pythoncom
import time
#import webbrowser
from threading import Timer
from dotenv import load_dotenv
import unicodedata

load_dotenv()

if getattr(sys, 'frozen', False):
        BASE_DIR = Path(sys._MEIPASS)
        USER_DIR = Path(os.path.dirname(sys.executable))
else:
        BASE_DIR = Path(__file__).resolve().parent
        USER_DIR = BASE_DIR              

        app = Flask(__name__)
        app.secret_key = os.getenv("FLASK_SECRET_KEY", "CLAVE_DE_RESPALDO_SECRETA")
        
        app.template_folder = str(BASE_DIR / "templates")
        app.static_folder = str(BASE_DIR / "static")
        
        BASE_DIR = Path(__file__).resolve().parent
        DB_PATH = BASE_DIR / "data" / "sistema.db"
        ARCHIVOS_DIR = BASE_DIR / "archivos"
        PLANTILLAS_DIR = ARCHIVOS_DIR / "plantillas"
        PLANT_CERT_DIR = PLANTILLAS_DIR / "certificados"
        PLANT_CONST_DIR = PLANTILLAS_DIR / "constancias"
        
        ACTIVIDADES_POR_CARRERA = {

        "INGENIERÍA CIVIL": (
            "● Elaboración de expedientes individuales con fines de tasación y diagnósticos técnicos legales.\n"
            "● Elaboración de memorias descriptivas, tasaciones.\n"
            "● Desarrollo de plano de viviendas, distribución, afectación, ubicación, perimétricos, entre otros.\n"
            "● Elaboración de certificados de búsqueda catastral, análisis técnico legal de predios afectados.\n"
            "● Elaboración de valorizaciones, estudios de mercado para viviendas, obras complementarias, cultivos, terrenos, entre otros.\n"
            "● Elaboración de mapas temáticos, conversión de archivos en formatos DWG, KMZ a SHP.\n"
            "● Recopilación y análisis de datos espaciales, imágenes raster, superposición de shapefiles, elaboración de planos clave, entre otros."
        ),
        "ARQUITECTURA": (
            "● Elaboración de expedientes individuales con fines de tasación y diagnósticos técnicos legales.\n"
            "● Elaboración de memorias descriptivas, tasaciones.\n"
            "● Desarrollo de plano de viviendas, distribución, afectación, ubicación, perimétricos, entre otros.\n"
            "● Elaboración de certificados de búsqueda catastral, análisis técnico legal de predios afectados.\n"
            "● Elaboración de valorizaciones, estudios de mercado para viviendas, obras complementarias, cultivos, terrenos, entre otros.\n"
            "● Elaboración de mapas temáticos, conversión de archivos en formatos DWG, KMZ a SHP.\n"
            "● Recopilación y análisis de datos espaciales, imágenes raster, superposición de shapefiles, elaboración de planos clave, entre otros."
        ),
        "INGENIERÍA GEOGRÁFICA": (
            "● Desarrollo de líneas base físicas, ambientales y biológicas\n"
            "● Análisis de base de datos alfanuméricos\n"
            "● Elaboración de mapas temáticos, conversión de archivos en formatos DWG, KMZ a SHP.\n"
            "● Recopilación y análisis de datos espaciales, imágenes raster, superposición de shapefiles, elaboración de planos clave, entre otros."
        ),
        "DISEÑO DE INTERIORES": (
            "● Elaboración de expedientes individuales con fines de tasación y diagnósticos técnicos legales.\n"
            "● Elaboración de memorias descriptivas, tasaciones.\n"
            "● Desarrollo de plano de viviendas, distribución, afectación, ubicación, perimétricos, entre otros.\n"
            "● Elaboración de certificados de búsqueda catastral, análisis técnico legal de predios afectados.\n"
            "● Elaboración de valorizaciones, estudios de mercado para viviendas, obras complementarias, cultivos, terrenos, entre otros.\n"
            "● Elaboración de mapas temáticos, conversión de archivos en formatos DWG, KMZ a SHP.\n"
            "● Recopilación y análisis de datos espaciales, imágenes raster, superposición de shapefiles, elaboración de planos clave, entre otros."
        ),
        "INGENIERÍA AMBIENTAL": (
            "● Elaboración de instrumentos de gestión ambiental.\n"
            "● Descripción de línea base física, biológica.\n"
            "● Levantamiento de observaciones y coordinaciones con la entidad.\n"
            "● Elaboración de matrices de impactos y planes de manejo ambiental.\n"
            "● Estudios de monitoreo ambiental (aire, agua, ruido, suelo) así como de monitoreos ocupacionales."
        ),
        "BIÓLOGO": (
            "● Elaboración de líneas base biológica\n"
            "● Descripción de flora y fauna presente en áreas de influencia para estudios ambientales\n"
            "● Participación en elaboración de instrumentos de gestión ambiental, referente al componente biológico\n"
            "● Desarrollo de monitoreos biológicos\n"
            "● Elaboración planes de muestreo biológicos"
        ),
        "SOCIÓLOGOANTROPÓLOGO": (
            "● Elaboración de líneas base social\n"
            "● Desarrollo de talleres y reuniones informativas generales\n"
            "● Elaboración de encuestas, toma de datos socioeconómicos entre otros\n"
            "● Análisis y elaboración de cuadros y gráficos estadísitcos e interpretación de los mismos\n"
            "● Elaboración de programas de ayuda social a poblaciones vulnerables"
        ),
        "INGENIERÍA AGRONOMÍAFORESTAL": (
            "● Elaboración de expedientes individuales con fines de tasación y diagnósticos técnicos legales.\n"
            "● Elaboración de memorias descriptivas, tasaciones.\n"
            "● Desarrollo de plano de viviendas, distribución, afectación, ubicación, perimétricos, entre otros.\n"
            "● Tasaciones de cultivos, árboles, obras complementarias y viviendas\n"
            "● Elaboración de memorias descriptivas y planos para expedientes individuales"
        ),
        "ADMINISTRACIÓN": (
            "● Elaboración y seguimientos de control del personal\n"
            "● Cálculo de pagos y liquidaciones a trabajadores\n"
            "● Evaluación de personal, entrevistas de trabajo, selección del personal\n"
            "● Envío de cartas y cotizaciones a clientes\n"
            "● Seguimiento de propuestas económicas, facturas en trámite\n"
            "● Apoyo en búsqueda de nuevos proyectos en la OSCE, armado de licitaciones y propuestas\n"
            "● Organización de archivos de trabajo, manuales operativos, documentación\n"
            "● Apoyo en el área técnicas para revisión de expedientes, fichas socioeconómicas, fichas técnicas de campo, revisión de documentos e informes entre otros.\n"
            "● Apoyo en la cotización de logística de campo, cronogramas, seguimiento de fechas de entregables, entre otros.\n"
            "● Asistencia en la gerencias"
        ),
        "DERECHO": (
            "● Análisis de partidas registrales, títulos archivados.\n"
            "● Revisión de expedientes técnicos legales, condición jurídica de afectados.\n"
            "● Elaboración de diagnósticos técnicos legales.\n"
            "● Análisis de predios en base al Decreto Legislativos N° 1192 y sus modificatorias.\n"
            "● Recopilación de documentación legal, constancias de posesión, partidas registrales, declaraciones juradas de colindantes.\n"
            "● Cuadro resumen de afectados, condición jurídica, entre otros.\n"
            "● Envío de cartas de solicitud de información registral y bases gráficas a entidades públicas y privadas."
        ),
        "MARKETING Y PUBLICIDAD": (
            "● Elaboración de banners, afiches\n"
            "● Manejo de redes sociales de la empresa (Facebook, Instagram, entre otros)\n"
            "● Ideas y propuestas de negocios para captar clientes en sector industria, transporte, minería e hidrocarburos, entre otros.\n"
            "● Obtención de base de datos de empresas a fin de contactarlos y captar potenciales clientes\n"
            "● Elaboración de brochure actualizado de la empresa, experiencia de la empresa, otros.\n"
            "● Asistencia en la gerencia\n"
            "● Seguimiento a clientes, inscripción y registro a entidades para elaboración de estudios ambientales\n"
            "● Posicionamiento de la marca a otros rubros\n"
            "● Comunity manager para redes sociales\n"
            "● Armado y elaboración de expedientes de licitación"
        )
    }
        MAP_CARRERAS = {
            "ARQUITECTURA Y URBANISMO": 1, "DERECHO": 2, "DIBUJANTE TECNICO MECANICO": 3, 
            "ING. INDUSTRIAL": 4, "ING. CIVIL": 5, "ADMINISTRACIÓN Y MARKETING": 6, 
            "ARQUITECTURA Y DISEÑO DE INT.": 7, "INGENIERÍA DE SISTEMAS": 8, 
            "ADMINISTRACION Y NEGOCIOS INTERNACIONALES": 9, "ADMINISTRACION DE EMPRESAS": 10, 
            "ECONÓMIA Y NEGOCIOS INTERNACIONALES": 11, "INGENIERIA DE SOFTWARE": 12, 
            "DISEÑO Y DESARROLLO DE MAQUINAS": 13, "DISEÑO Y ADMINISTRACIÓN BANCARIA Y FINANCIERA": 14, 
            "INGENIERÍA AMBIENTAL": 15, "INGENIERIA GEOLOGICA": 16, "INGENIERÍA EN GESTIÓN AMBIENTAL": 17, 
            "PSICOLOGIA": 18, "INGENIERIA GEOGRÀFICA": 19, "DISEÑO GRAFICO": 20, 
            "ING. EN SEGURIDAD LABORAL Y AMBIENTAL": 21, "INGENIERÍA COMERCIAL": 22, 
            "INGENIERÍA BIOTECNOLÓGICA": 23, "INGENIERÍA AGRÍCOLA": 24, "SOCIOLOGIA": 25, 
            "INGENIERÍA DE CIBERSEGURIDAD": 26, "SIN DEFINIR..": 27, "ARQUITECTURA": 28, 
            "INGENIERA INDUSTRIAL Y DE SISTEMAS": 29, "INGENIERÍA DE SISTEMAS DE  INFORMACION": 30, 
            "ADMINISTRACIÓN": 31, "PUBLICIDAD Y MULTIMEDIA": 33, "ADMINISTRACIÓN HOTELERA Y TURISMO": 34, 
            "INGENIERIA DE GESTION EMPRESARIAL": 35, "CONTABILIDAD": 36, "ARQUITECTURA Y DISEÑO": 37, 
            "ADMINISTRACION DE HOTELES Y TURISMO": 38, "DESARROLLO DE SOFTWARE": 51, 
            "INGENIERIA MECANICA": 53, "INGENIERÍA DE MINAS": 55, "GEOGRAFÍA Y MEDIO AMBIENTE": 56, 
            "BIOINGENIERIA": 57, "BIOLOGÍA": 58
        }

        MAP_INSTITUCIONES = {
            "SIN DEFINIR": 1, "UNIVERSIDAD PRIVADA DEL NORTE": 2, "SENATI": 3, 
            "UNIVERSIDAD CESAR VALLEJO": 4, "UNIVERSIDAD PERUANA DE CIENCIAS APLICADAS": 5, 
            "UNIVERSIDAD DE LIMA": 6, "UNIVERSISAD SAN MARTIN DE PORRES": 7,
            "UNIVERSIDAD SAN MARTIN DE PORRES": 7,
            "UNIVERSIDAD NACIONAL DE SAN CRISTOBAL DE HUAMANGA": 8, "UNIVERSIDAD NACIONAL DE CAJAMARCA": 9, 
            "UNIVERSIDAD CATÓLICA SAN PABLO": 10, "UNIVERSIDAD DE PIURA": 11, 
            "UNIVERSIDAD NACIONAL MAYOR DE SAN MARCOS": 12, "UNIVERSIDAD NACIONAL DE INGENIERIA": 13, 
            "UNIVERSIDAD RICARDO PALMA": 14, "UNIVERSIDAD CATOLICA SANTO TORIBIO DE MOGROVEJO": 15, 
            "UNIVERSIDAD NACIONAL DE SAN AGUSTÍN DE AREQUIPA": 16, "UNIVERSIDAD CATÓLICA LOS ÁNGELES DE CHIMBOTE": 17, 
            "UNIVERSIDAD PRIVADA ANTENOR ORREGO": 19, "UNIVERSIDAD AUTÓNOMA DEL PERÚ": 20, 
            "UNIVERSIDAD NACIONAL DE MOQUEGUA": 21, "UNIVERSIDAD CATOLICA DE SANTA MARIA": 22, 
            "UNIVERSIDAD NACIONAL DE SAN ANTONIO ABAD DEL CUSCO": 23, "UNIVERSIDAD CONTINENTAL": 24, 
            "UNIVERSIDAD TECNOLOGICA DEL PERU": 25, "IDAT": 26, 
            "UNIVERSIDAD NACIONAL FEDERICO VILLAREAL (UNFV)": 27, "UNIVERSIDAD NACIONAL TECNOLÓGICA DE LIMA SUR": 29, 
            "UNIVERSIDAD NACIONAL DE SAN MARTÍN": 30, "UNIVERSIDAD SAN IGNACIO DE LOYOLA": 31, 
            "UNIVERSIDAD FEMENINA DEL SAGRADO CORAZÓN (UNIFE)": 32, "PONTIFICIA UNIVERSIDAD CATÓLICA DEL PERÚ": 33, 
            "UNIVERSIDAD NACIONAL DE PIURA": 34, "UNIVERSIDAD CIENTÍFICA DEL SUR": 35, 
            "UNIVERSIDAD NACIONAL PEDRO RUIZ GALLO": 36, "UNIVERSIDAD DE INGENIERÍA Y TECNOLOGÍA (UTEC)": 37, 
            "UNIVERSIDAD PRIVADA SAN JUAN BAUTISTA": 38, "UNIVERSIDAD NACIONAL DEL CALLAO": 39, 
            "UNIVERSIDAD NACIONAL AGRARIA LA MOLINA": 40, "UNIVERSIDAD NACIONAL DE FRONTERA": 41, 
            "UNIVERSIDAD DE CIENCIAS Y ARTES DE AMÉRICA LATINA (UCAL)": 42, "ESCUELA DE ADMINISTRACIÓN DE NEGOCIOS PARA GRADUADOS (ESAN)": 43, 
            "UNIVERSIDAD ANDINA DEL CUSCO": 44, "UNIVERSIDAD TECNOLOGICA DE LOS ANDES": 45, 
            "INSTITUTO DE EDUCACIÓN SUPERIOR PRIVADO ZEGEL": 46, "UNIVERSIDAD NACIONAL JOSÉ FAUSTINO SÁNCHEZ CARRIÓN (UNJFSC)": 47, 
            "UNIVERSIDAD PERUANA UNIÓN": 49, "UNIVERSIDAD CATOLICA SEDES SAPIENTIAE (UCSS)": 50, 
            "UNIVERSIDAD DEL PACÍFICO": 51, "UNIVERSIDAD NACIONAL DE EDUCACIÓN ENRIQUE GUZMÁN Y VALLE": 52, 
            "UNIVERSIDAD NACIONAL SANTIAGO ANTÚNEZ DE MAYOLO": 53, "UNIVERSIDAD NACIONAL INTERCULTURAL DE LA SELVA CENTRAL JUAN SANTOS ATAHUALPA": 54, 
            "UNIVERSIDAD NACIONAL DE LA AMAZONIA PERUANA": 55, "UNIVERSIDAD NACIONAL INTERCULTURAL DE LA AMAZONIA": 56, 
            "TOULOUSE LAUTREC": 57, "UNIVERSIDAD NACIONAL DE BARRANCA": 58, 
            "UNIVERSIDAD NACIONAL DE TRUJILLO": 59, "UNIVERSIDAD NACIONAL DE UCAYALI": 60, 
            "UNIVERSIDAD NACIONAL DEL SANTA": 65, "COLEGIO INGENIERÍA SAC": 66, 
            "I.E.P. SANTA ROSA": 67, "ACADEMIA PRE-U": 68, "CENTRO DE IDIOMAS AMERICANO": 72, 
            "I.E.P. ADVENTISTA": 73, "INSTITUTO TECNOLÓGICO SUR": 74, "UNIVERSIDAD DE HUANUCO": 102
        }

        MAP_FACULTADES = {
            "POR DEFINIR...": 1, "INGENIERÍA": 2, "DERECHO": 3, "CIENCIAS CONTABLES": 4, 
            "ADMINISTRACIÓN": 5, "DISEÑO": 6, "LICENCIATURA": 7, "ARQUITECTURA": 8, 
            "INGENIERÍA Y ARQUITECTURA": 9, "CIENCIAS ECONÓMICAS Y ADMINISTRATIVA": 10, 
            "INGENIERÍA CIVIL": 11, "INGENERÍA CIVIL Y ARQUITECTURA": 12, "ARQUITECTURA URBANISMO Y ARTES": 13, 
            "INGENIERÍA DE SISTEMAS E INFORMÁTICA": 14, "CIENCIAS SOCIALES": 15, "CIENCIAS E INGENIERÍA": 16, 
            "INGENIERIA AMBIENTAL": 17, "CIENCIAS EMPRESARIALES": 18, "ARQUITECTURA E INGENIERIA": 19, 
            "ARQUITECTURA Y DISEÑO": 20, "ING. GEOLÓGICA, MINERA, METALÚRGICA Y GEOGRÁFICA": 22, 
            "FACULTAD DE PROCESOS": 23, "FACULTAD DE LETRAS Y CIENCIAS HUMANAS": 26, 
            "INGENIERIA QUIMICA": 27, "NEGOCIOS": 28
        }




ALLOWED_TEMPLATE_EXT = {".docx"}

    # URL CSV publicada
SHEETS_CSV_URL = "https://docs.google.com/spreadsheets/d/e/2PACX-1vRjpQ8p9e9e-95_0tLHiC1gMKc6GMdeEHvJb3gvQUyStrbRxQXfv97Fi6XwS__sRCHaOWqpMHpu_48i/pub?output=csv"


MYSQL_HOST = 'mysql-a21bb78-sistemasnet26-321c.k.aivencloud.com'       
MYSQL_USER = 'avnadmin'           
MYSQL_PASSWORD = os.getenv("DB_PASSWORD") 
MYSQL_DB = 'SistemaGenerador'
MYSQL_PORT = 10658


def db_mysql():
    return pymysql.connect(
        host=MYSQL_HOST,
        user=MYSQL_USER,
        password=MYSQL_PASSWORD,
        database=MYSQL_DB,
        port=MYSQL_PORT,
        ssl={"ssl": {}},
        cursorclass=pymysql.cursors.DictCursor 
    )

def db_mysql_gmingenieros():
    return pymysql.connect(
        host=MYSQL_HOST,
        user=MYSQL_USER,
        password=MYSQL_PASSWORD,
        database='gmingenieros',
        port=MYSQL_PORT,
        ssl={"ssl": {}},
        cursorclass=pymysql.cursors.DictCursor 
    )

def db():
    DB_PATH.parent.mkdir(parents=True, exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON;")
    return conn


def ahora():
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")

def buscar_id(valor_excel, diccionario):
    if not valor_excel: return None
    if valor_excel in diccionario:
        return diccionario[valor_excel]
    for key, val in diccionario.items():
        if key in valor_excel or valor_excel in key:
            return val
    return None

def normalizar_texto(texto):
    if not texto: return ""
    texto = str(texto).upper().strip()
    texto_sin_tildes = ''.join(c for c in unicodedata.normalize('NFD', texto) if unicodedata.category(c) != 'Mn')
    return texto_sin_tildes

def buscar_id(valor_excel, diccionario):
    if not valor_excel: return None
    
    val_norm = normalizar_texto(valor_excel)
    
   
    for key, val in diccionario.items():
        if normalizar_texto(key) == val_norm:
            return val
            

    for key, val in diccionario.items():
        key_norm = normalizar_texto(key)
        if val_norm in key_norm or key_norm in val_norm:
            return val
            
    return None

def login_required(view_func):
        def wrapper(*args, **kwargs):
            if not session.get("user_id"):
                flash("Inicia sesión para continuar", "info")
                return redirect(url_for("login"))
            return view_func(*args, **kwargs)
        wrapper.__name__ = view_func.__name__
        return wrapper


def usuario_actual():
        if not session.get("user_id"):
            return None
        conn = db_mysql()
        u = conn.execute("SELECT * FROM usuarios WHERE id = ?", (session["user_id"],)).fetchone()
        conn.close()
        return u

def formatear_fecha_latam(valor):
    txt = s(valor)
    if not txt:
        return ""
    
    formatos = ["%d/%m/%Y", "%d/%m/%y", "%m/%d/%Y", "%Y-%m-%d"]
    for fmt in formatos:
        try:
            dt = datetime.strptime(txt, fmt)
            return dt.strftime("%d/%m/%Y") 
        except ValueError:
            continue
            
    return txt

def s(v):
    if pd.isna(v) or v is None:
        return ""
    txt = str(v).strip()
    
    if txt.lower() == "nan":
        return ""
        
    if txt.endswith(".0"):
        txt = txt[:-2]
        
    return txt

def upper(v):
        return s(v).upper()

def parse_form_datetime_to_iso(value):
        # Espera "d/m/yyyy HH:MM:SS" o "dd/mm/yyyy HH:MM:SS"
        txt = s(value)
        if not txt:
            return ""
        for fmt in ("%d/%m/%Y %H:%M:%S", "%d/%m/%Y %H:%M"):
            try:
                dt = datetime.strptime(txt, fmt)
                return dt.strftime("%Y-%m-%d %H:%M:%S")
            except ValueError:
                pass
        return ""


def ensure_solicitudes_schema(conn):
        cur = conn.cursor()

        cur.execute("""
        CREATE TABLE IF NOT EXISTS solicitudes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            sheet_uid TEXT UNIQUE,
            marca_temporal TEXT,
            marca_dt TEXT,
            correo_solicitante TEXT,
            correo TEXT,
            tipo_documento TEXT,
            nombres TEXT,
            apellidos TEXT,
            documento TEXT,
            fecha_inicio TEXT,
            fecha_fin TEXT,
            universidad TEXT,
            codigo_alumno TEXT,
            facultad TEXT,
            carrera TEXT,
            ciclo TEXT,
            cargo TEXT,
            actividades TEXT,
            horas_totales INTEGER,
            estado TEXT,
            observaciones TEXT,
            fecha_revision TEXT,
            revisado_por TEXT,
            fecha_emision TEXT,
            emitido_por TEXT,
            codigo_documento TEXT,
            ruta_pdf TEXT,
            creado_en TEXT,
            actualizado_en TEXT
        )
        """)

        # 2 Detecta columnas existentes
        existentes = {r["name"] for r in conn.execute("PRAGMA table_info(solicitudes)").fetchall()}

        # 3 Agrega columnas faltantes de manera segura
        # Nota importante, en SQLite no puedes añadir NOT NULL sin DEFAULT
        columnas = {
            "sheet_uid": "TEXT",
            "marca_temporal": "TEXT",
            "marca_dt": "TEXT",
            "correo_solicitante": "TEXT",
            "correo": "TEXT",
            "tipo_documento": "TEXT",
            "nombres": "TEXT",
            "apellidos": "TEXT",
            "documento": "TEXT",
            "fecha_inicio": "TEXT",
            "fecha_fin": "TEXT",
            "universidad": "TEXT",
            "codigo_alumno": "TEXT",
            "facultad": "TEXT",
            "carrera": "TEXT",
            "ciclo": "TEXT",
            "cargo": "TEXT",
            "actividades": "TEXT",
            "horas_totales": "INTEGER",
            "estado": "TEXT",
            "observaciones": "TEXT",
            "fecha_revision": "TEXT",
            "revisado_por": "TEXT",
            "fecha_emision": "TEXT",
            "emitido_por": "TEXT",
            "codigo_documento": "TEXT",
            "ruta_pdf": "TEXT",
            "creado_en": "TEXT",
            "actualizado_en": "TEXT",
        }

        for col, tipo in columnas.items():
            if col not in existentes:
                cur.execute(f"ALTER TABLE solicitudes ADD COLUMN {col} {tipo}")

        # 4 Backfill para que no te vuelvan a salir errores por valores vacíos
        # Copia correo_solicitante a correo si correo está vacío
        cur.execute("""
            UPDATE solicitudes
            SET correo = correo_solicitante
            WHERE (correo IS NULL OR TRIM(correo) = '')
            AND (correo_solicitante IS NOT NULL AND TRIM(correo_solicitante) <> '')
        """)

        # Si actualizado_en está vacío, lo igualamos a creado_en
        cur.execute("""
            UPDATE solicitudes
            SET actualizado_en = creado_en
            WHERE (actualizado_en IS NULL OR TRIM(actualizado_en) = '')
            AND (creado_en IS NOT NULL AND TRIM(creado_en) <> '')
        """)

        # 5 Índices
        cur.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_solicitudes_sheet_uid ON solicitudes(sheet_uid)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_solicitudes_estado ON solicitudes(estado)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_solicitudes_documento ON solicitudes(documento)")
        cur.execute("CREATE INDEX IF NOT EXISTS idx_solicitudes_marca_dt ON solicitudes(marca_dt)")

        conn.commit()


def ensure_config_schema(conn):
        cur = conn.cursor()
        cur.execute("""
        CREATE TABLE IF NOT EXISTS configuracion (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            ruta_salida TEXT NOT NULL DEFAULT 'archivos/emitidos/',
            correo_emisor TEXT NOT NULL DEFAULT '',
            envio_correo INTEGER NOT NULL DEFAULT 0,
            actualizado_en TEXT NOT NULL
        )
        """)

        row = cur.execute("SELECT id FROM configuracion WHERE id = 1").fetchone()
        if not row:
            cur.execute("""
                INSERT INTO configuracion (id, ruta_salida, correo_emisor, envio_correo, actualizado_en)
                VALUES (1, 'archivos/emitidos/', '', 0, ?)
            """, (ahora(),))
        conn.commit()


def get_config(conn):
        ensure_config_schema(conn)
        return conn.execute("SELECT * FROM configuracion WHERE id = 1").fetchone()


@app.get("/")
def root():
        if session.get("user_id"):
            return redirect(url_for("dashboard"))
        return redirect(url_for("login"))


@app.get("/login")
def login():
        return render_template("login.html")


@app.post("/iniciar-sesion")
def iniciar_sesion():
    correo = (request.form.get("correo") or "").strip().lower()
    contrasena = (request.form.get("contrasena") or "").strip()

    if not correo or not contrasena:
        flash("Completa correo y contraseña", "error")
        return redirect(url_for("login"))

    conn = db_mysql()
    with conn.cursor() as cur:
        cur.execute("SELECT * FROM usuarios WHERE correo = %s", (correo,))
        u = cur.fetchone()

        if not u:
            conn.close()
            flash("Credenciales incorrectas", "error")
            return redirect(url_for("login"))

        if int(u["activo"]) != 1:
            conn.close()
            flash("Esperar confirmacion del coordinador", "error")
            return redirect(url_for("login"))

        if not check_password_hash(u["password_hash"], contrasena):
            conn.close()
            flash("Credenciales incorrectas", "error")
            return redirect(url_for("login"))

        cur.execute("UPDATE usuarios SET ultimo_acceso = %s WHERE id = %s", (ahora(), u["id"]))
        conn.commit()
    conn.close()

    session["user_id"] = u["id"]
    session["rol"] = u["rol"]
    session["nombre"] = u["nombre_completo"]

    flash("Sesión iniciada", "success")
    return redirect(url_for("dashboard"))

@app.get("/registro")
def registro():
    return render_template("registro.html")

@app.post("/registrarse")
def registrarse():
    nombre = (request.form.get("nombre") or "").strip()
    correo = (request.form.get("correo") or "").strip().lower()
    contrasena = (request.form.get("contrasena") or "").strip()
    confirmar = (request.form.get("confirmar") or "").strip()

    if not nombre or not correo or not contrasena or not confirmar:
        flash("Completa todos los campos", "error")
        return redirect(url_for("registro"))

    if contrasena != confirmar:
        flash("Las contraseñas no coinciden", "error")
        return redirect(url_for("registro"))

    conn = db_mysql()
    with conn.cursor() as cur:
        cur.execute("SELECT 1 FROM usuarios WHERE correo = %s", (correo,))
        existe = cur.fetchone()
        
        if existe:
            conn.close()
            flash("Ese correo ya está registrado", "error")
            return redirect(url_for("registro"))

        cur.execute("SELECT COUNT(*) AS n FROM usuarios")
        total = cur.fetchone()["n"]
        
        if total == 0:
            rol = "COORDINADOR"
            activo = 1  
        else:
            rol = "ASISTENTE"
            activo = 0  

        password_hash = generate_password_hash(contrasena)

        cur.execute("""
            INSERT INTO usuarios (nombre_completo, correo, password_hash, rol, activo, fecha_creacion)
            VALUES (%s, %s, %s, %s, %s, %s)
        """, (nombre, correo, password_hash, rol, activo, ahora()))
        conn.commit()
    conn.close()


    if activo == 1:
        flash("Registro exitoso. Eres el primer usuario, ingresas como COORDINADOR.", "success")
    else:
        flash("Registro exitoso. Tu cuenta está inactiva, esperar confirmación del coordinador.", "info")
        
    return redirect(url_for("login"))


@app.get("/logout")
def logout():
        session.clear()
        flash("Sesión cerrada", "info")
        return redirect(url_for("login"))


@app.get("/dashboard")
@login_required
def dashboard():
    nombre_usuario_actual = session.get("nombre", "")

    # 1. KPIs
    conn_lite = db()
    def count_estado_usuario(estado):
        if session.get("rol") == "COORDINADOR":
            r = conn_lite.execute("SELECT COUNT(*) AS n FROM solicitudes WHERE estado = ?", (estado,)).fetchone()
        else:
            if estado == "EMITIDO":
                r = conn_lite.execute("SELECT COUNT(*) AS n FROM solicitudes WHERE estado = ? AND emitido_por = ?", (estado, nombre_usuario_actual)).fetchone()
            elif estado == "REVISADO":
                r = conn_lite.execute("SELECT COUNT(*) AS n FROM solicitudes WHERE estado = ? AND revisado_por = ?", (estado, nombre_usuario_actual)).fetchone()
            else:
                r = conn_lite.execute("SELECT COUNT(*) AS n FROM solicitudes WHERE estado = ?", (estado,)).fetchone()
        return int(r["n"]) if r else 0

    kpi_recibido  = count_estado_usuario("RECIBIDO")
    kpi_pendiente = count_estado_usuario("PENDIENTE")
    kpi_observado = count_estado_usuario("OBSERVADO")
    kpi_revisado  = count_estado_usuario("REVISADO")
    kpi_emitido   = count_estado_usuario("EMITIDO")
    kpi_anulado   = count_estado_usuario("ANULADO")
    conn_lite.close()

    # 2. Historial de Actividades: DIRECTO DE MYSQL
    conn_my = db_mysql()
    ultimas10 = []
    try:
        with conn_my.cursor() as cur:
            if session.get("rol") == "COORDINADOR":
                cur.execute("""
                    SELECT id, ficha AS dni, usuario, accion AS estado, creado_en AS fecha, nombre_completo, tipo_documento
                    FROM historial_solicitud ORDER BY id DESC LIMIT 10
                """)
            else:
                cur.execute("""
                    SELECT id, ficha AS dni, usuario, accion AS estado, creado_en AS fecha, nombre_completo, tipo_documento
                    FROM historial_solicitud WHERE usuario = %s ORDER BY id DESC LIMIT 10
                """, (nombre_usuario_actual,))
            
            ultimas10 = cur.fetchall()
    finally:
        conn_my.close()

    return render_template(
        "dashboard.html", active="dashboard",
        kpi_recibido=kpi_recibido, kpi_pendiente=kpi_pendiente, kpi_observado=kpi_observado,
        kpi_revisado=kpi_revisado, kpi_emitido=kpi_emitido, kpi_anulado=kpi_anulado,
        ultimas10=ultimas10, alert_observadas=kpi_observado, alert_plantillas=0, alert_correos_fallidos=0
    )

def add_historial(solicitud_id, ficha, usuario, accion, detalle=None, nombre_comp="N/A", tipo_doc="DOC"):
    try:
        conn = db_mysql()
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO historial_solicitud (solicitud_id, ficha, usuario, accion, detalle, creado_en, nombre_completo, tipo_documento)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
            """, (solicitud_id, ficha, usuario, accion, detalle, ahora(), nombre_comp, tipo_doc))
        conn.commit()
        conn.close()
    except Exception as e:
        print(f"Error guardando historial en MySQL: {e}")


@app.get("/solicitudes")
@login_required
def solicitudes():
        estado = (request.args.get("estado") or "").strip().upper()
        tipo = (request.args.get("tipo") or "").strip().upper()
        dni = (request.args.get("dni") or "").strip()
        desde = (request.args.get("desde") or "").strip()  # YYYY-MM-DD
        hasta = (request.args.get("hasta") or "").strip()  # YYYY-MM-DD

        conn = db()

        where = []
        params = []

        if estado:
            where.append("estado = ?")
            params.append(estado)

        if tipo:
            where.append("tipo_documento = ?")
            params.append(tipo)

        if dni:
            where.append("documento LIKE ?")
            params.append(f"%{dni}%")

        # filtro por fecha usando marca_dt
        # desde y hasta vienen como YYYY-MM-DD
        if desde:
            where.append("(marca_dt >= ? OR (marca_dt IS NULL AND creado_en >= ?))")
            params.append(desde + " 00:00:00")
            params.append(desde + " 00:00:00")

        if hasta:
            where.append("(marca_dt <= ? OR (marca_dt IS NULL AND creado_en <= ?))")
            params.append(hasta + " 23:59:59")
            params.append(hasta + " 23:59:59")

        sql = "SELECT * FROM solicitudes"
        if where:
            sql += " WHERE " + " AND ".join(where)
        sql += " ORDER BY id DESC LIMIT 500"

        rows = conn.execute(sql, params).fetchall()
        total = conn.execute("SELECT COUNT(*) AS n FROM solicitudes").fetchone()["n"]

        conn.close()

        return render_template(
            "solicitudes.html",
            active="solicitudes",
            solicitudes=rows,
            total=total,
            estados=["RECIBIDO","PENDIENTE","OBSERVADO","REVISADO","EMITIDO","ANULADO"],
            filtros={"estado": estado, "tipo": tipo, "dni": dni, "desde": desde, "hasta": hasta}
        )

@app.get("/solicitudes/<int:sid>")
@login_required
def solicitudes_detalle(sid):
        conn = db()

        s = get_solicitud_por_id(conn, sid)
        conn.close() 

        if not s:
            flash("Solicitud no encontrada", "error")
            return redirect(url_for("solicitudes"))

       
        historial = get_historial(sid)

        return render_template(
            "detalle.html",
            active="solicitudes",
            s=s,
            historial=historial
        )
def get_solicitud_por_id(conn, sid):
        return conn.execute("SELECT * FROM solicitudes WHERE id = ?", (sid,)).fetchone()

def get_historial(sid):
        conn_my = db_mysql()
        try:
            with conn_my.cursor() as cur:
                cur.execute("""
                    SELECT * FROM historial_solicitud
                    WHERE solicitud_id = %s
                    ORDER BY id DESC
                """, (sid,))
                return cur.fetchall()
        except Exception as e:
            print(f"Error leyendo historial en MySQL: {e}")
            return []
        finally:
            conn_my.close()

@app.post("/solicitudes/<int:sid>/guardar")
@login_required
def solicitudes_guardar(sid):
    horas = (request.form.get("horas_totales") or "").strip()
    observ = (request.form.get("observaciones") or "").strip()
    conn = db()
    s = get_solicitud_por_id(conn, sid)
    if not s:
        conn.close(); flash("Solicitud no encontrada", "error"); return redirect(url_for("solicitudes"))

    cambios = []
    if (s["horas_totales"] or "") != horas: cambios.append(f"Horas {s['horas_totales'] or ''} -> {horas}")
    if (s["observaciones"] or "") != observ: cambios.append("Observaciones actualizadas")

    conn.execute("UPDATE solicitudes SET horas_totales = ?, observaciones = ?, actualizado_en = ? WHERE id = ?", (horas, observ, ahora(), sid))
    conn.commit()
    conn.close()

    if cambios:
        # AHORA PASAMOS EL NOMBRE
        add_historial(sid, s["documento"] or "", session.get("nombre", "USUARIO"), "GUARDAR", " | ".join(cambios), f"{s['nombres']} {s['apellidos']}", s["tipo_documento"])

    flash("Cambios guardados", "success"); return redirect(url_for("solicitudes_detalle", sid=sid))


@app.post("/solicitudes/<int:sid>/estado/observado")
@login_required
def solicitudes_marcar_observado(sid):
    conn = db(); s = get_solicitud_por_id(conn, sid)
    if not s: conn.close(); flash("Solicitud no encontrada", "error"); return redirect(url_for("solicitudes"))
    conn.execute("UPDATE solicitudes SET estado = 'OBSERVADO', actualizado_en = ? WHERE id = ?", (ahora(), sid))
    conn.commit(); conn.close()

    add_historial(sid, s["documento"] or "", session.get("nombre","USUARIO"), "OBSERVADO", "Marcado como OBSERVADO", f"{s['nombres']} {s['apellidos']}", s["tipo_documento"])
    flash("Estado actualizado a OBSERVADO", "success"); return redirect(url_for("solicitudes_detalle", sid=sid))

@app.post("/solicitudes/<int:sid>/estado/revisado")
@login_required
def solicitudes_marcar_revisado(sid):
    conn = db(); s = get_solicitud_por_id(conn, sid)
    if not s: conn.close(); flash("Solicitud no encontrada", "error"); return redirect(url_for("solicitudes"))
    conn.execute("UPDATE solicitudes SET estado = 'REVISADO', revisado_por = ?, fecha_revision = ?, actualizado_en = ? WHERE id = ?", (session.get("nombre","USUARIO"), ahora(), ahora(), sid))
    conn.commit(); conn.close()

    add_historial(sid, s["documento"] or "", session.get("nombre","USUARIO"), "REVISADO", "Marcado como REVISADO", f"{s['nombres']} {s['apellidos']}", s["tipo_documento"])
    flash("Estado actualizado a REVISADO", "success"); return redirect(url_for("solicitudes_detalle", sid=sid))

@app.route("/solicitudes/<int:sid>/emitir", methods=['GET', 'POST'])
@login_required
def solicitudes_emitir(sid):
        conn = db()
        config = get_config(conn) 

        s = get_solicitud_por_id(conn, sid)
        if not s:
            conn.close()
            flash("Solicitud no encontrada", "error")
            return redirect(url_for("solicitudes"))

        horas = str(s["horas_totales"] or "").strip()
        if not horas:
            conn.close()
            flash("Antes de emitir, registra las horas totales", "error")
            return redirect(url_for("solicitudes_detalle", sid=sid))

        tipo_doc = s["tipo_documento"]
        carrera = s["carrera"]
        plantilla = None
        try:
            conn_my = db_mysql()
            with conn_my.cursor() as cur_my:
                cur_my.execute("""
                    SELECT ruta_docx FROM plantillas 
                    WHERE tipo_documento = %s AND carrera = %s AND activo = 1
                """, (tipo_doc, carrera))
                plantilla = cur_my.fetchone()
            conn_my.close()
        except Exception as err_tpl:
            print(f"Error buscando plantilla en MySQL: {err_tpl}")

        if not plantilla:
            conn.close()
            flash(f"No hay plantilla activa para {tipo_doc} de {carrera}", "error")
            return redirect(url_for("solicitudes_detalle", sid=sid))

        ruta_plantilla_abs = (BASE_DIR / plantilla["ruta_docx"]).resolve()
        
        try:
            doc = Document(ruta_plantilla_abs)
            
            
            reemplazos = {
                "{{NOMBRE_COMPLETO}}": f"{s['nombres']} {s['apellidos']}",
                "{{DNI}}": str(s["documento"] or ""),
                "{{UNIVERSIDAD}}": str(s["universidad"] or ""),
                "{{FACULTAD}}": str(s["facultad"] or ""),
                "{{CARRERA}}": str(s["carrera"] or ""),
                "{{CODIGO}}": str(s["codigo_alumno"] or ""),
                "{{FECHA_INICIO}}": formatear_fecha_latam(s["fecha_inicio"]), 
                "{{FECHA_FIN}}": formatear_fecha_latam(s["fecha_fin"]),
                "{{CARGO}}": str(s["cargo"] or "ASISTENTE").upper(),
                "{{HORAS_TOTALES}}": str(horas),
                "{{ACTIVIDADES}}": str(s["actividades"] or ""),
                "{{FECHA_EMISION}}": datetime.now().strftime("%d/%m/%Y")
            }

        
            def reemplazar_texto_en_parrafo(parrafo, reemplazos):
                texto_parrafo = parrafo.text
                for llave, valor in reemplazos.items():
                    if llave in texto_parrafo:
                        texto_parrafo = texto_parrafo.replace(llave, valor)
                
            
                if parrafo.text != texto_parrafo:
                
                    parrafo.clear()
                    parrafo.add_run(texto_parrafo)

        
            for parrafo in doc.paragraphs:
                reemplazar_texto_en_parrafo(parrafo, reemplazos)
                    
            
            for tabla in doc.tables:
                for fila in tabla.rows:
                    for celda in fila.cells:
                        for parrafo in celda.paragraphs:
                            reemplazar_texto_en_parrafo(parrafo, reemplazos)
            
            nombre_archivo = f"{tipo_doc}_{s['documento']}_{uuid.uuid4().hex[:6]}.docx"
            ruta_salida_dir = BASE_DIR / config["ruta_salida"]
            ruta_salida_dir.mkdir(parents=True, exist_ok=True)
            
            ruta_archivo_final = ruta_salida_dir / nombre_archivo
            doc.save(ruta_archivo_final)
            
            ruta_bd = str(ruta_archivo_final.relative_to(BASE_DIR)).replace("\\", "/")
            
        except Exception as e:
            conn.close()
            flash(f"Error al generar el documento: {e}", "error")
            return redirect(url_for("solicitudes_detalle", sid=sid))
            
        codigo_doc = f"{tipo_doc}-{s['documento']}-{datetime.now().strftime('%y%m%d')}"

        conn.execute("""
            UPDATE solicitudes
            SET estado = 'EMITIDO', emitido_por = ?, fecha_emision = ?, actualizado_en = ?, codigo_documento = ?, ruta_pdf = ?
            WHERE id = ?
        """, (session.get("nombre","USUARIO"), ahora(), ahora(), codigo_doc, ruta_bd, sid))
        conn.commit()
        conn.close()

     
        add_historial(sid, s["documento"] or "", session.get("nombre","USUARIO"), "EMISION", "Documento generado y EMITIDO", f"{s['nombres']} {s['apellidos']}", s["tipo_documento"])

        try:
            conn_my = db_mysql()
            with conn_my.cursor() as cur:
             
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS reportes (
                        solicitud_id INT PRIMARY KEY,
                        codigo_documento VARCHAR(100),
                        nombre_completo VARCHAR(255),
                        documento VARCHAR(50),
                        tipo_documento VARCHAR(50),
                        estado VARCHAR(50),
                        fecha_emision DATETIME,
                        emitido_por VARCHAR(100),
                        ruta_pdf VARCHAR(500)
                    )
                """)
                
                nombre_completo = f"{s['nombres']} {s['apellidos']}"
                cur.execute("""
                    INSERT INTO reportes (solicitud_id, codigo_documento, nombre_completo, documento, tipo_documento, estado, fecha_emision, emitido_por, ruta_pdf)
                    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s)
                    ON DUPLICATE KEY UPDATE 
                        estado = 'EMITIDO', fecha_emision = VALUES(fecha_emision), emitido_por = VALUES(emitido_por), ruta_pdf = VALUES(ruta_pdf)
                """, (
                    sid, codigo_doc, nombre_completo, s['documento'], tipo_doc, 
                    'EMITIDO', ahora(), session.get("nombre","USUARIO"), ruta_bd
                ))
            conn_my.commit()
            conn_my.close()
        except Exception as e:
          
            flash(f"Advertencia: Documento generado, pero hubo un error en MySQL: {e}", "warning")
            print(f"Error al guardar reporte en MySQL: {e}")

        flash("Solicitud marcada como EMITIDO y documento generado", "success")
        return redirect(url_for("solicitudes_detalle", sid=sid))

       

@app.post("/solicitudes/<int:sid>/anular")
@login_required
def solicitudes_anular(sid):
    conn = db(); s = get_solicitud_por_id(conn, sid)
    if not s: conn.close(); flash("Solicitud no encontrada", "error"); return redirect(url_for("solicitudes"))
    conn.execute("UPDATE solicitudes SET estado = 'ANULADO', actualizado_en = ? WHERE id = ?", (ahora(), sid))
    conn.commit(); conn.close()

    add_historial(sid, s["documento"] or "", session.get("nombre","USUARIO"), "ANULADO", "Solicitud ANULADA", f"{s['nombres']} {s['apellidos']}", s["tipo_documento"])

    try:
        conn_my = db_mysql()
        with conn_my.cursor() as cur:
            cur.execute("UPDATE reportes SET estado = 'ANULADO' WHERE solicitud_id = %s", (sid,))
        conn_my.commit(); conn_my.close()
    except Exception as e: print(f"Error al anular reporte en MySQL: {e}")

    flash("Solicitud anulada", "success"); return redirect(url_for("solicitudes_detalle", sid=sid))


@app.post("/solicitudes/sincronizar")
@login_required
def solicitudes_sincronizar():
    try:
        url_con_cache_breaker = f"{SHEETS_CSV_URL}&cache_buster={int(time.time())}"
        df = pd.read_csv(url_con_cache_breaker)
        df.columns = df.columns.str.strip()
    except Exception as e:
        flash(f"Error al conectar con Google Sheets: {e}", "error")
        return redirect(url_for("solicitudes"))

    conn = db()
    ensure_solicitudes_schema(conn)
    cur = conn.cursor()

    try:
        conn_gm = db_mysql_gmingenieros()
        cur_gm = conn_gm.cursor()
    except Exception as e:
        conn.close()
        flash(f"Error al conectar con BD gmingenieros: {e}", "error")
        return redirect(url_for("solicitudes"))


    nuevos = duplicados = errores = omitidos = 0

    for _, row in df.iterrows():
        marca = s(row.get("Marca temporal"))
        documento = s(row.get("N° DOCUMENTO"))

        if not documento or not marca:
            continue

        existe = cur.execute("SELECT 1 FROM solicitudes WHERE documento = ?", (documento,)).fetchone()
        if existe:
            duplicados += 1
            continue 

        estado_inicial = 'RECIBIDO' 
        try:
            cur_gm.execute("SELECT 1 FROM empleados WHERE dni = %s", (documento,))
            if cur_gm.fetchone():
                estado_inicial = 'IGNORADO'
                omitidos += 1
        except Exception as e:
            print(e)
        # -------------------

        tipo_raw = upper(row.get("Seleccione lo que desea solicitar"))
        nombres = upper(row.get("NOMBRES"))
        apellidos = upper(row.get("APELLIDOS"))
        f_inicio = formatear_fecha_latam(row.get("Fecha de inicio (dd/mm/yyyy)"))
        f_fin = formatear_fecha_latam(row.get("Fecha de fin (dd/mm/yyyy)"))
        uni = upper(row.get("NOMBRE DE LA UNIVERSIDAD O INSTITUTO"))
        cod_alumno = s(row.get("CODIGO DE ALUMNO"))
        facultad = upper(row.get("FACULTAD"))
        carrera_excel = upper(row.get("CARRERA")) 
        ciclo = s(row.get("CICLO"))
        cargo = upper(row.get("CARGO"))
        
        actividades_final = "" 
        encontrado = False
        
        for carrera_mapeada, texto_actividades in ACTIVIDADES_POR_CARRERA.items():
            if carrera_mapeada in carrera_excel or carrera_excel in carrera_mapeada:
                actividades_final = texto_actividades
                encontrado = True
                break
        
        if not encontrado:
            actividades_final = s(row.get("ACTIVIDADES")) 

        correo = s(row.get("CORREO ELECTRONICO") or row.get("CORREO ELECTRÓNICO") or row.get("Dirección de correo electrónico") or "")

        if "CONSTANCIA" in tipo_raw:
            tipo = "CONST"
        elif "PRE-PROFESIONALES" in tipo_raw or "PRE PROFESIONALES" in tipo_raw:
            tipo = "CERT_PRAC_PRE"
        elif "PROFESIONALES" in tipo_raw:
            tipo = "CERT_PRAC_PROF"
        elif "TRABAJO" in tipo_raw:
            tipo = "CERT_TRAB"
        elif "ACEPTACION" in tipo_raw:
            tipo = "CART_ACEPT"
        elif "RECOMENDACION" in tipo_raw:
            tipo = "CART_RECOM"
        elif "RECONOCIMIENTO" in tipo_raw:
            tipo = "CERT_RECON"
        else:
            tipo = "CERT"

        sheet_uid = f"{marca}|{documento}|{tipo}|{int(time.time())}"

        try:
            
            cur.execute("""
                INSERT INTO solicitudes (
                    sheet_uid, marca_temporal, correo, correo_solicitante, tipo_documento,
                    nombres, apellidos, documento,
                    fecha_inicio, fecha_fin,
                    universidad, codigo_alumno, facultad, carrera, ciclo, cargo,
                    actividades, estado, creado_en, actualizado_en
                )
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (
                sheet_uid, marca, correo, correo, tipo,
                nombres, apellidos, documento,
                f_inicio, f_fin,
                uni, cod_alumno, facultad, carrera_excel, ciclo, cargo,
                actividades_final, estado_inicial, ahora(), ahora() 
            ))
           
            if estado_inicial == 'RECIBIDO':
                nuevos += 1
                
                try:
                    
                    id_uni = buscar_id(uni, MAP_INSTITUCIONES)
                    id_facultad = buscar_id(facultad, MAP_FACULTADES)
                    id_carrera = buscar_id(carrera_excel, MAP_CARRERAS)

                    # Calculamos el próximo IdEmpleado
                    cur_gm.execute("SELECT IFNULL(MAX(IdEmpleado), 0) + 1 AS next_id FROM empleados")
                    next_id = cur_gm.fetchone()["next_id"]
                    
                
                    cur_gm.execute("""
                        INSERT INTO empleados (
                            IdEmpleado, Dni, Nombres, Apellidos, Correo, 
                            CodigoEstudiante, Ciclo, estado,
                            IdInstitucionEducativa, IdFacultad, IdCarrera
                        )
                        VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                    """, (
                        next_id, documento, nombres, apellidos, correo, 
                        cod_alumno, ciclo, '1',
                        id_uni, id_facultad, id_carrera
                    ))
                except Exception as error_mysql:
                    print(f"Error al guardar en la tabla empleados (MySQL): {error_mysql}")

        except Exception as e:
            print(f"Error general de inserción: {e}") 
            errores += 1

    conn_gm.commit() 
    conn_gm.close()
    
    conn.commit()   
    conn.close()

    registrados_bd = nuevos + omitidos 
    
    if errores > 0:
        flash(f"Sincronización: {nuevos} nuevos, {omitidos} ignorados, {registrados_bd} registrados en la bd. {errores} errores.", "warning")
    else:
        flash(f"Sincronización exitosa: {nuevos} nuevos, {omitidos} ignorados, {duplicados} duplicados, {registrados_bd} registrados en la bd.", "success")
        
    return redirect(url_for("solicitudes"))


@app.get("/plantillas")
@login_required
def plantillas():
        conn = db_mysql()
        with conn.cursor() as cur:
            cur.execute("""
                SELECT *
                FROM plantillas
                ORDER BY actualizado_en DESC, id DESC
            """)
            rows = cur.fetchall()
            
            cur.execute("SELECT COUNT(*) AS n FROM plantillas")
            total_row = cur.fetchone()
            total = total_row["n"] if total_row else 0
        conn.close()

        return render_template(
            "plantillas.html",
            active="plantillas",
            plantillas=rows,
            total_plantillas=total
        )

@app.post("/plantillas/subir")
@login_required
def plantillas_subir():
        tipo_input = (request.form.get("tipo_documento") or "").strip().upper()
        carrera = (request.form.get("carrera") or "").strip().upper()
        f = request.files.get("archivo")

        tipos_validos = {"CERT_TRAB", "CERT_PRAC_PROF", "CERT_PRAC_PRE", "CART_ACEPT", "CART_RECOM", "CERT_RECON", "CERT", "CONST"}
        if tipo_input not in tipos_validos:
            flash(f"Tipo inválido: {tipo_input}", "error") 
            return redirect(url_for("plantillas"))

        tipo_db = tipo_input

        if not carrera:
            flash("Carrera es obligatoria", "error")
            return redirect(url_for("plantillas"))

        if not f or not f.filename:
            flash("Debes seleccionar un archivo docx", "error")
            return redirect(url_for("plantillas"))

        filename = secure_filename(f.filename)
        ext = Path(filename).suffix.lower()
        if ext not in ALLOWED_TEMPLATE_EXT:
            flash("Solo se permite .docx", "error")
            return redirect(url_for("plantillas"))

        dest_dir = PLANT_CERT_DIR if tipo_db == "CERT" else PLANT_CONST_DIR
        dest_dir.mkdir(parents=True, exist_ok=True)

        safe_carrera = "_".join(carrera.split())
        final_name = f"{tipo_input}_{safe_carrera}.docx" 
        dest_path = dest_dir / final_name

        try:
            f.save(dest_path)
        except Exception as e:
            flash(f"No se pudo guardar el archivo: {e}", "error")
            return redirect(url_for("plantillas"))

        ahora_txt = ahora()
        ruta_rel = str(dest_path.relative_to(BASE_DIR)).replace("\\", "/")
        
        # Conexión y guardado en MySQL
        conn = db_mysql()
        with conn.cursor() as cur:
            cur.execute("""
                INSERT INTO plantillas (tipo_documento, carrera, archivo_nombre, ruta_docx, activo, creado_en, actualizado_en)
                VALUES (%s, %s, %s, %s, 1, %s, %s)
                ON DUPLICATE KEY UPDATE 
                    archivo_nombre = VALUES(archivo_nombre), 
                    ruta_docx = VALUES(ruta_docx), 
                    activo = 1, 
                    actualizado_en = VALUES(actualizado_en)
            """, (tipo_db, carrera, final_name, ruta_rel, ahora_txt, ahora_txt))
        conn.commit()
        conn.close()
        
        flash("Plantilla procesada y guardada correctamente", "success")
        return redirect(url_for("plantillas"))

@app.get("/usuarios")
@login_required
def usuarios():
    q = (request.args.get("q") or "").strip().lower()
    rol = (request.args.get("rol") or "").strip().upper()
    activo = (request.args.get("activo") or "").strip()

    conn = db_mysql()
    with conn.cursor() as cur:
        where = []
        params = []

        if q:
            where.append("(LOWER(nombre_completo) LIKE %s OR LOWER(correo) LIKE %s)")
            params.extend([f"%{q}%", f"%{q}%"])
        if rol:
            where.append("rol = %s")
            params.append(rol)
        if activo in ("0", "1"):
            where.append("activo = %s")
            params.append(int(activo))

        sql = "SELECT id, nombre_completo, correo, rol, activo, ultimo_acceso, fecha_creacion FROM usuarios"
        if where:
            sql += " WHERE " + " AND ".join(where)
        sql += " ORDER BY id DESC LIMIT 500"

        cur.execute(sql, params)
        rows = cur.fetchall()
    conn.close()

    return render_template(
        "usuarios.html", active="usuarios", usuarios=rows,
        filtros={"q": q, "rol": rol, "activo": activo},
        roles=["COORDINADOR", "ASISTENTE"]
    )

@app.post("/usuarios/crear")
@login_required
def usuarios_crear():
    if (session.get("rol") or "").upper() != "COORDINADOR":
        flash("No tienes permiso para crear usuarios", "error")
        return redirect(url_for("usuarios"))

    nombre = (request.form.get("nombre") or "").strip()
    correo = (request.form.get("correo") or "").strip().lower()
    rol = (request.form.get("rol") or "").strip().upper()
    contrasena = (request.form.get("contrasena") or "").strip()

    if not nombre or not correo or not rol or not contrasena:
        flash("Completa todos los campos", "error")
        return redirect(url_for("usuarios"))
    if rol not in ("COORDINADOR", "ASISTENTE"):
        flash("Rol inválido", "error")
        return redirect(url_for("usuarios"))

    conn = db_mysql()
    with conn.cursor() as cur:
        cur.execute("SELECT 1 FROM usuarios WHERE correo = %s", (correo,))
        if cur.fetchone():
            conn.close()
            flash("Ese correo ya está registrado", "error")
            return redirect(url_for("usuarios"))

        password_hash = generate_password_hash(contrasena)
        cur.execute("""
            INSERT INTO usuarios (nombre_completo, correo, password_hash, rol, activo, fecha_creacion, ultimo_acceso)
            VALUES (%s, %s, %s, %s, 1, %s, NULL)
        """, (nombre, correo, password_hash, rol, ahora()))
        conn.commit()
    conn.close()

    flash("Usuario creado", "success")
    return redirect(url_for("usuarios"))

@app.post("/usuarios/<int:uid>/toggle")
@login_required
def usuarios_toggle(uid):
    if (session.get("rol") or "").upper() != "COORDINADOR":
        flash("No tienes permiso", "error")
        return redirect(url_for("usuarios"))
    if session.get("user_id") == uid:
        flash("No puedes modificar tu propio usuario", "error")
        return redirect(url_for("usuarios"))

    conn = db_mysql()
    with conn.cursor() as cur:
        cur.execute("SELECT id, activo FROM usuarios WHERE id = %s", (uid,))
        u = cur.fetchone()
        if not u:
            conn.close()
            flash("Usuario no encontrado", "error")
            return redirect(url_for("usuarios"))

        nuevo = 0 if int(u["activo"]) == 1 else 1
        cur.execute("UPDATE usuarios SET activo = %s WHERE id = %s", (nuevo, uid))
        conn.commit()
    conn.close()

    flash("Estado actualizado", "success")
    return redirect(url_for("usuarios"))

@app.post("/usuarios/<int:user_id>/cambiar-rol")
@login_required
def usuarios_cambiar_rol(user_id):
    if (session.get("rol") or "").upper() != "COORDINADOR":
        flash("No tienes permiso", "error")
        return redirect(url_for("usuarios"))

    nuevo_rol = (request.form.get("rol") or "").strip().upper()
    if nuevo_rol not in ("COORDINADOR", "ASISTENTE"):
        flash("Rol inválido", "error")
        return redirect(url_for("usuarios"))

    conn = db_mysql()
    with conn.cursor() as cur:
        cur.execute("SELECT id FROM usuarios WHERE id = %s", (user_id,))
        if not cur.fetchone():
            conn.close()
            flash("Usuario no encontrado", "error")
            return redirect(url_for("usuarios"))

        cur.execute("UPDATE usuarios SET rol = %s WHERE id = %s", (nuevo_rol, user_id))
        conn.commit()
    conn.close()

    flash("Rol actualizado", "success")
    return redirect(url_for("usuarios"))

@app.post("/usuarios/<int:user_id>/eliminar")
@login_required
def usuarios_eliminar(user_id):
    if (session.get("rol") or "").upper() != "COORDINADOR":
        flash("No tienes permiso", "error")
        return redirect(url_for("usuarios"))
    if session.get("user_id") == user_id:
        flash("No puedes eliminarte a ti mismo", "error")
        return redirect(url_for("usuarios"))

    conn = db_mysql()
    with conn.cursor() as cur:
        cur.execute("DELETE FROM usuarios WHERE id = %s", (user_id,))
        conn.commit()
    conn.close()

    flash("Usuario eliminado", "success")
    return redirect(url_for("usuarios"))

@app.post("/plantillas/<int:pid>/toggle")
@login_required
def plantillas_toggle(pid):
        conn = db_mysql()
        with conn.cursor() as cur:
            cur.execute("SELECT id, activo FROM plantillas WHERE id = %s", (pid,))
            row = cur.fetchone()
            
            if not row:
                conn.close()
                flash("Plantilla no encontrada", "error")
                return redirect(url_for("plantillas"))

            nuevo = 0 if int(row["activo"]) == 1 else 1

            cur.execute("""
                UPDATE plantillas
                SET activo = %s, actualizado_en = %s
                WHERE id = %s
            """, (nuevo, ahora(), pid))
        conn.commit()
        conn.close()

        flash("Estado actualizado", "success")
        return redirect(url_for("plantillas"))

@app.get("/plantillas/<int:pid>/descargar")
@login_required
def plantillas_descargar(pid):
        conn = db_mysql()

        p = conn.execute("SELECT * FROM plantillas WHERE id = %s", (pid,)).fetchone()
        conn.close()

        if not p:
            flash("Plantilla no encontrada", "error")
            return redirect(url_for("plantillas"))

        ruta_rel = p["ruta_docx"]
        ruta_abs = (BASE_DIR / ruta_rel).resolve()

        if not str(ruta_abs).startswith(str(BASE_DIR.resolve())):
            abort(403)

        if not ruta_abs.exists():
            flash("El archivo no existe en disco", "error")
            return redirect(url_for("plantillas"))

        return send_file(ruta_abs, as_attachment=True, download_name=p["archivo_nombre"])

@app.get("/reportes")
@login_required
def reportes():
        tipo = (request.args.get("tipo") or "").strip().upper()
        estado = (request.args.get("estado") or "").strip().upper()
        desde = (request.args.get("desde") or "").strip()
        hasta = (request.args.get("hasta") or "").strip()

        where = []
        params = []
        
        if tipo:
            where.append("tipo_documento = %s")
            params.append(tipo)
        if estado:
            where.append("estado = %s")
            params.append(estado)
        if desde:
            where.append("fecha_emision >= %s")
            params.append(desde + " 00:00:00")
        if hasta:
            where.append("fecha_emision <= %s")
            params.append(hasta + " 23:59:59")

        sql = """
            SELECT 
                solicitud_id AS id,
                codigo_documento AS codigo,
                nombre_completo,
                documento,
                tipo_documento AS tipo,
                estado,
                fecha_emision,
                emitido_por,
                ruta_pdf
            FROM reportes
        """
        
        where.append("estado IN ('EMITIDO', 'ANULADO')")

        if where:
            sql += " WHERE " + " AND ".join(where)
        sql += " ORDER BY fecha_emision DESC LIMIT 500"

   
        conn = db_mysql()
        with conn.cursor() as cur:
            cur.execute(sql, params)
            rows = cur.fetchall()
        conn.close()

        total = len(rows)

        return render_template(
            "reportes.html",
            active="reportes",
            reportes=rows,
            total=total,
            filtros={"tipo": tipo, "estado": estado, "desde": desde, "hasta": hasta},
        )

@app.get("/configuracion")
@login_required
def configuracion():

        conn_my = db_mysql()
        with conn_my.cursor() as cur:
            cur.execute("SELECT * FROM usuarios WHERE id = %s", (session["user_id"],))
            u = cur.fetchone()
        conn_my.close()

        conn_lite = db()
        config = get_config(conn_lite)
        conn_lite.close()

        return render_template(
            "configuracion.html",
            active="configuracion",
            usuario=u,
            config=config
        )

@app.post("/configuracion/guardar")
@login_required
def configuracion_guardar():
        ruta_salida = (request.form.get("ruta_salida") or "").strip()
        correo_emisor = (request.form.get("correo_emisor") or "").strip()
        envio_correo = 1 if request.form.get("envio_correo") == "1" else 0

        if not ruta_salida:
            flash("La ruta de salida no puede estar vacía", "error")
            return redirect(url_for("configuracion"))

        # normaliza slash final
        ruta_salida = ruta_salida.replace("\\", "/")
        if not ruta_salida.endswith("/"):
            ruta_salida += "/"

        conn = db()
        ensure_config_schema(conn)
        conn.execute("""
            UPDATE configuracion
            SET ruta_salida = ?, correo_emisor = ?, envio_correo = ?, actualizado_en = ?
            WHERE id = 1
        """, (ruta_salida, correo_emisor, envio_correo, ahora()))
        conn.commit()
        conn.close()

        flash("Configuración guardada", "success")
        return redirect(url_for("configuracion"))

@app.post("/configuracion/cambiar-password")
@login_required
def configuracion_cambiar_password():
        actual = (request.form.get("password_actual") or "").strip()
        nueva = (request.form.get("password_nueva") or "").strip()
        confirmar = (request.form.get("password_confirmar") or "").strip()

        if not actual or not nueva or not confirmar:
            flash("Completa todos los campos de contraseña", "error")
            return redirect(url_for("configuracion"))

        if nueva != confirmar:
            flash("La nueva contraseña y su confirmación no coinciden", "error")
            return redirect(url_for("configuracion"))

        if len(nueva) < 6:
            flash("La nueva contraseña debe tener al menos 6 caracteres", "error")
            return redirect(url_for("configuracion"))

        conn = db()
        u = conn.execute("SELECT * FROM usuarios WHERE id = ?", (session["user_id"],)).fetchone()

        if not u or not check_password_hash(u["password_hash"], actual):
            conn.close()
            flash("La contraseña actual es incorrecta", "error")
            return redirect(url_for("configuracion"))

        conn.execute("""
            UPDATE usuarios
            SET password_hash = ?
            WHERE id = ?
        """, (generate_password_hash(nueva), session["user_id"]))
        conn.commit()
        conn.close()

        flash("Contraseña actualizada", "success")
        return redirect(url_for("configuracion"))

@app.get("/documento/<int:doc_id>/ver")
@login_required
def ver_pdf(doc_id):
        conn = db()
        s = get_solicitud_por_id(conn, doc_id)
        conn.close()
        
        if not s or not s["ruta_pdf"]:
            flash("Documento no encontrado", "error")
            return redirect(url_for("reportes"))
            
        ruta_abs = (BASE_DIR / s["ruta_pdf"]).resolve()
        return send_file(ruta_abs, mimetype='application/pdf', as_attachment=False)

@app.get("/documento/<int:doc_id>/descargar")
@login_required
def descargar_doc(doc_id):
        conn = db()
        s = get_solicitud_por_id(conn, doc_id)
        conn.close()
        
        if not s or not s["ruta_pdf"]:
            flash("Documento no encontrado", "error")
            return redirect(url_for("reportes"))
            
        ruta_abs = (BASE_DIR / s["ruta_pdf"]).resolve()
        nombre_descarga = Path(s["ruta_pdf"]).name
        return send_file(ruta_abs, as_attachment=True, download_name=nombre_descarga)

def open_browser():
        webbrowser.open_new("http://127.0.0.1:5000/")


with app.app_context():
    connection = db()
    ensure_solicitudes_schema(connection) 
    ensure_config_schema(connection)
    connection.close() 

if __name__ == "__main__":
        Timer(1.5, open_browser).start()
        app.run(port=5000, debug=False) 