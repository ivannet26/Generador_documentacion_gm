import sqlite3
from docx import Document
from datetime import datetime

PLANTILLA = "CERTIFICADO DE PRÁCTICAS_Alvaro Martinez.docx.docx"
SALIDA = "certificado_generado.docx"
dni = "76348612"  # cambia por uno real

# fecha formal
meses = {
    1:"enero",2:"febrero",3:"marzo",4:"abril",
    5:"mayo",6:"junio",7:"julio",8:"agosto",
    9:"septiembre",10:"octubre",11:"noviembre",12:"diciembre"
}
hoy = datetime.now()
fecha_formal = f"{hoy.day} de {meses[hoy.month]} del {hoy.year}"

# Traer de BD solo lo que podamos
conn = sqlite3.connect("certificados.db")
cur = conn.cursor()

cur.execute("""
SELECT nombres, apellidos, documento, carrera, universidad, codigo_alumno, fecha_inicio, fecha_fin
FROM solicitudes
WHERE documento = ?
""", (dni,))
row = cur.fetchone()
conn.close()

if not row:
    print("No encontrado en BD")
    raise SystemExit

nombres, apellidos, documento, carrera, universidad, codigo_alumno, fecha_inicio, fecha_fin = row

# Diccionario de marcadores según tu plantilla
datos = {
    "{{NOMBRE_COMPLETO}}": f"{nombres} {apellidos}".upper(),
    "{{DNI}}": str(documento).strip(),
    "{{CARRERA}}": (str(carrera).upper() if carrera else ""),
    "{{UNIVERSIDAD}}": (str(universidad).upper() if universidad else ""),
    "{{CODIGO}}": (str(codigo_alumno).strip() if codigo_alumno else ""),
    "{{FECHA_INICIO}}": (str(fecha_inicio).strip() if fecha_inicio else ""),
    "{{FECHA_FIN}}": (str(fecha_fin).strip() if fecha_fin else ""),
    "{{FECHA_EMISION}}": fecha_formal,
}

doc = Document(PLANTILLA)

# Reemplazo en párrafos
for p in doc.paragraphs:
    for k, v in datos.items():
        if k in p.text:
            for run in p.runs:
                run.text = run.text.replace(k, v)

# Reemplazo dentro de tablas si tu plantilla tuviera
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for k, v in datos.items():
                    if k in p.text:
                        for run in p.runs:
                            run.text = run.text.replace(k, v)

doc.save(SALIDA)
print("Generado:", SALIDA)
